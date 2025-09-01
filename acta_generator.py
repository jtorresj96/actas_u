# -*- coding: utf-8 -*-
"""
acta_generator
--------------
Genera un ACTA (.docx) a partir del contenido de una transcripción (.docx),
usando LLM para construir:
- Resumen ejecutivo
- Orden del día
- Desarrollo del orden del día
- Decisiones tomadas
- Compromisos

Requisitos:
    pip install -U python-docx openai

Uso:
    from openai import OpenAI
    from acta_generator import build_acta_from_transcription
    client = OpenAI()  # requiere OPENAI_API_KEY en el entorno
    acta_path, sections = build_acta_from_transcription(
        transcripcion_docx_path=".../transcripcion.docx",
        acta_file_path=".../ACTA.docx",
        client=client,
        model="gpt-5"  # o el que uses
    )
"""

from __future__ import annotations
import os, re
from typing import Optional, Dict, Tuple
from docx import Document
from docx.shared import Pt
from openai import OpenAI


# -------- Helpers de formato / IO --------

def _extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    return "\n".join(parts)


def _add_paragraphs(doc: Document, text: str, heading_level: Optional[int] = None) -> None:
    """
    Aplica reglas simples de formato:
    - Si heading_level se pasa, se crea un heading y se retorna.
    - Líneas que empiezan por "n. " -> Heading 3
    - Viñetas que empiezan por -, • o * -> estilo 'List Bullet'
    - "Campo: valor" -> Campo en negrita
    - Resto -> párrafos normales tamaño 11pt
    """
    if not text:
        return

    if heading_level:
        doc.add_heading(text.strip(), level=heading_level)
        return

    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue

        # Títulos numerados → Heading 3
        if re.match(r"^\d+\.\s", line):
            doc.add_heading(line, level=3)
            continue

        # Viñetas
        if re.match(r"^[-•*]\s", line):
            p = doc.add_paragraph(line, style='List Bullet')
            for run in p.runs:
                run.font.size = Pt(11)
            continue

        # Campos tipo Responsable: ...
        m = re.match(r"^([\wÁÉÍÓÚÜÑáéíóúüñ\s]+:\s)(.*)$", line)
        if m:
            p = doc.add_paragraph()
            campo, valor = m.group(1), m.group(2)
            r1 = p.add_run(campo)
            r1.bold = True
            r1.font.size = Pt(11)
            r2 = p.add_run(valor)
            r2.font.size = Pt(11)
            continue

        # Párrafo normal
        p = doc.add_paragraph(line)
        # Asegurar tamaño 11pt
        if not p.runs:
            p.add_run(line)
        for run in p.runs:
            run.font.size = Pt(11)


# -------- LLM wrappers --------

def _chat_complete(client: OpenAI, model: str, prompt: str) -> str:
    """Wrapper simple para chat.completions."""
    resp = client.chat.completions.create(
        model=model,
        messages=[{"role": "user", "content": prompt}],
    )
    return (resp.choices[0].message.content or "").strip()


# -------- Prompts base --------

def _build_prompts(contenido_transcripcion: str, orden_dia_text: Optional[str] = None) -> Dict[str, str]:
    prompt_resumen = (
        "Redacta un resumen ejecutivo profesional y claro sobre la siguiente reunión. "
        "No uses formato markdown. Incluye: (1) temas tratados, (2) decisiones clave, (3) acuerdos relevantes. "
        "Hazlo entendible para un lector externo que no estuvo presente.\n\n"
        f"Transcripción:\n{contenido_transcripcion}"
    )

    prompt_orden_dia = (
        "Identifica los puntos tratados como Orden del Día. Escribe solo los temas numerados, una línea por tema, "
        "con el formato '1. Tema', '2. Tema', etc. "
        "Evita introducciones y no uses formato markdown.\n\n"
        f"Transcripción:\n{contenido_transcripcion}"
    )

    if orden_dia_text is None:
        orden_dia_text = "{orden_dia}"  # placeholder si no se pasó aún

    prompt_desarrollo = (
        "Desarrolla cada punto del orden del día en tres partes: (1) descripción general, "
        "(2) resumen de la discusión, (3) implicaciones o acuerdos. Formatea como:\n\n"
        "1. Título del punto\nDescripción: ...\nResumen: ...\nImplicaciones: ...\n\n"
        "No uses markdown ni frases introductorias. Hazlo claro y formal.\n\n"
        f"Orden del Día:\n{orden_dia_text}\n\n"
        f"Transcripción:\n{contenido_transcripcion}"
    )

    prompt_decisiones = (
        "Extrae decisiones claras tomadas en la reunión. Presenta cada decisión como una viñeta que empiece con '- '. "
        "No uses markdown. Sé concreto y profesional.\n\n"
        f"Transcripción:\n{contenido_transcripcion}"
    )

    prompt_compromisos = (
        "Extrae compromisos asumidos por los asistentes. Por cada uno incluye: Responsable, Tarea, Fecha (si aplica). "
        "Evita markdown, presenta cada compromiso claramente en párrafos separados.\n\n"
        f"Transcripción:\n{contenido_transcripcion}"
    )

    return {
        "resumen": prompt_resumen,
        "orden_dia": prompt_orden_dia,
        "desarrollo": prompt_desarrollo,
        "decisiones": prompt_decisiones,
        "compromisos": prompt_compromisos,
    }


# -------- Función pública principal --------

def build_acta_from_transcription(
    transcripcion_docx_path: str,
    acta_file_path: str,
    client: Optional[OpenAI] = None,
    model: str = "gpt-5",
) -> Tuple[str, Dict[str, str]]:
    """
    Lee el DOCX de transcripción, genera secciones con LLM y escribe el ACTA final (.docx).

    Returns:
        (acta_file_path, sections_dict)
        sections_dict = {
            "resumen": str,
            "orden_dia": str,
            "desarrollo": str,
            "decisiones": str,
            "compromisos": str
        }
    """
    if client is None:
        # Requiere OPENAI_API_KEY en el entorno
        client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

    contenido_transcripcion = _extract_text_from_docx(transcripcion_docx_path)

    # Construir prompts base
    prompts = _build_prompts(contenido_transcripcion)

    # 1) Resumen y Orden del Día
    contenido_resumen   = _chat_complete(client, model, prompts["resumen"])
    contenido_orden_dia = _chat_complete(client, model, prompts["orden_dia"])

    # 2) Desarrollo (inyectando el Orden del Día generado)
    prompt_desarrollo_final = _build_prompts(contenido_transcripcion, orden_dia_text=contenido_orden_dia)["desarrollo"]
    contenido_desarrollo = _chat_complete(client, model, prompt_desarrollo_final)

    # 3) Decisiones y Compromisos
    contenido_decisiones  = _chat_complete(client, model, prompts["decisiones"])
    contenido_compromisos = _chat_complete(client, model, prompts["compromisos"])

    # --- Escribir ACTA ---
    os.makedirs(os.path.dirname(acta_file_path), exist_ok=True)
    doc = Document()

    # Título principal
    doc.add_heading("ACTA DE LA REUNIÓN", level=1)

    # Secciones
    _add_paragraphs(doc, "Resumen Ejecutivo", heading_level=2)
    _add_paragraphs(doc, contenido_resumen)

    _add_paragraphs(doc, "Orden del Día", heading_level=2)
    _add_paragraphs(doc, contenido_orden_dia)

    _add_paragraphs(doc, "Desarrollo del Orden del Día", heading_level=2)
    _add_paragraphs(doc, contenido_desarrollo)

    _add_paragraphs(doc, "Decisiones Tomadas", heading_level=2)
    _add_paragraphs(doc, contenido_decisiones)

    _add_paragraphs(doc, "Compromisos y Próxima Reunión", heading_level=2)
    _add_paragraphs(doc, contenido_compromisos)

    doc.save(acta_file_path)

    sections = {
        "resumen": contenido_resumen,
        "orden_dia": contenido_orden_dia,
        "desarrollo": contenido_desarrollo,
        "decisiones": contenido_decisiones,
        "compromisos": contenido_compromisos,
    }

    return acta_file_path, sections
