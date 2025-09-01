# -*- coding: utf-8 -*-
"""
openai_transcription
--------------------
Transcribe audios largos a DOCX sin timestamps, en párrafos largos (~15 s),
dividiendo el audio en chunks con solape para evitar pérdidas.

Requisitos:
    pip install -U openai pydub python-docx tqdm

Notas:
- Por defecto usa el modelo "whisper-1". Puedes cambiarlo a otro si lo prefieres.
- pydub requiere ffmpeg instalado en el sistema.
"""

from __future__ import annotations
import os
import unicodedata
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Iterator, List, Optional, Tuple, Union

from pydub import AudioSegment
from docx import Document
from tqdm.auto import tqdm
from openai import OpenAI
from acta_generator import build_acta_from_transcription


# ==========================
# Helpers genéricos
# ==========================

def _norm_nfc(p: Union[str, Path]) -> str:
    """Normaliza rutas con tildes/ñ para evitar problemas en distintos FS."""
    return str(Path(unicodedata.normalize("NFC", str(p))))

def _fmt_ts(seconds: Optional[float]) -> str:
    if seconds is None:
        return "00:00:00.000"
    total = int(seconds)
    ms = int(round((seconds - total) * 1000))
    h = total // 3600
    m = (total % 3600) // 60
    s = total % 60
    return f"{h:02d}:{m:02d}:{s:02d}.{ms:03d}"

def _seg_val(seg, name, default=None):
    """Lee atributos de segmentos tanto si son objeto como dict."""
    if hasattr(seg, name):
        return getattr(seg, name)
    if isinstance(seg, dict):
        return seg.get(name, default)
    return default

def _merge_text_no_dupe(current: str, new: str, max_overlap_words: int = 12) -> str:
    """Evita duplicar texto cuando hay solape entre chunks."""
    if not current:
        return (new or "").strip()
    if not new:
        return current
    a = current.strip().split()
    b = new.strip().split()
    max_k = min(max_overlap_words, len(a), len(b))
    for k in range(max_k, 2, -1):
        if a[-k:] == b[:k]:
            return current.rstrip() + " " + " ".join(b[k:])
    return current.rstrip() + " " + new.lstrip()


# ==========================
# Configuración de transcripción
# ==========================

@dataclass
class TranscriptionConfig:
    # Chunks
    chunk_minutes: int = 10
    overlap_seconds: int = 2

    # Audio preproc
    amplify_db: int = 6

    # Salida de párrafos
    target_block_sec: float = 15.0
    max_slack_sec: float = 7.0  # 15–22 s aprox
    include_timestamps: bool = False  # Tu pedido: False

    # Export temporal a MP3 para evitar 413
    export_mp3: bool = True
    mp3_bitrate: str = "64k"
    mp3_ar: str = "16000"
    mp3_ac: str = "1"  # mono

    # OpenAI
    model: str = "whisper-1"
    force_language: Optional[str] = "es"  # None = autodetect
    openai_api_key: Optional[str] = None  # si None, usa env var OPENAI_API_KEY

    # UI
    show_progress: bool = True
    save_incremental: bool = True  # guarda el docx cada vez que cierra un bloque


# ==========================
# Núcleo de procesamiento
# ==========================

def _iter_chunks(audio: AudioSegment, chunk_ms: int, overlap_ms: int) -> Iterator[Tuple[int, int, AudioSegment]]:
    """Genera (start_ms, end_ms, segmento) con solape."""
    duration_ms = len(audio)
    start = 0
    while start < duration_ms:
        end = min(start + chunk_ms, duration_ms)
        yield start, end, audio[start:end]
        if end >= duration_ms:
            break
        start = end - overlap_ms

def _write_block(doc: Document, text: str, include_ts: bool = False,
                 start_ts: Optional[float] = None, end_ts: Optional[float] = None) -> None:
    """Escribe un párrafo en el DOCX; por defecto sin timestamps."""
    text = (text or "").strip()
    if not text:
        return
    if include_ts and start_ts is not None and end_ts is not None:
        doc.add_paragraph(f"[{_fmt_ts(start_ts)} – {_fmt_ts(end_ts)}] {text}")
    else:
        doc.add_paragraph(text)




# ==========================
# API principal (pública)
# ==========================

def transcribe_to_docx(
    audio_file: Union[str, Path],
    output_docx: Union[str, Path],
    config: Optional[TranscriptionConfig] = None,
) -> Tuple[str, int]:
    """
    Transcribe un audio largo a DOCX, en párrafos de ~15 s sin timestamps.

    Args:
        audio_file: ruta del .wav/.mp3 de entrada.
        output_docx: ruta del .docx de salida.
        config: TranscriptionConfig opcional (valores por defecto si None).

    Returns:
        (ruta_docx, bloques_escritos)
    """
    cfg = config or TranscriptionConfig()

    # --- OpenAI client ---
    api_key = cfg.openai_api_key or os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("Falta OPENAI_API_KEY (pásala en config.openai_api_key o como env var).")
    client = OpenAI(api_key=api_key)

    # --- Carga y preproc de audio ---
    audio_path = _norm_nfc(audio_file)
    print("Esto es el audio path: ", audio_path)
    if not Path(audio_path).exists():
        raise FileNotFoundError(f"No se encontró el audio: {audio_path}")

    audio = AudioSegment.from_file(audio_path)
    if cfg.amplify_db:
        audio = audio + cfg.amplify_db

    duration_ms = len(audio)
    chunk_ms = int(cfg.chunk_minutes * 60 * 1000)
    overlap_ms = int(cfg.overlap_seconds * 1000)

    # --- Prepara DOCX ---
    output_docx = _norm_nfc(output_docx)
    Path(output_docx).parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading(f"Transcripción", level=1)
    doc.add_paragraph(f"Archivo: {Path(audio_path).name}")
    duration = _fmt_ts(duration_ms/1000.0)
    doc.add_paragraph(f"Duración total: {duration}")
    doc.add_paragraph(
        f"Párrafos objetivo ≈ {int(cfg.target_block_sec)} s "
        f"(tolerancia {int(cfg.max_slack_sec)} s)"
    )
    doc.add_heading("Transcripción (párrafos largos ~15 s)", level=2)
    doc.save(output_docx)

    # --- Estado de bloques ---
    state = {
        "current_text": "",
        "block_start": None,  # segundos globales
        "last_end": None,     # segundos globales
        "blocks_count": 0,
    }

    def flush_block():
        if state["current_text"]:
            _write_block(
                doc,
                state["current_text"],
                include_ts=cfg.include_timestamps,
                start_ts=state["block_start"],
                end_ts=state["last_end"],
            )
            if cfg.save_incremental:
                doc.save(output_docx)
            state["blocks_count"] += 1
        state["current_text"] = ""
        state["block_start"] = None
        state["last_end"] = None

    # --- Iteración de chunks ---
    chunks = list(_iter_chunks(audio, chunk_ms, overlap_ms))
    progress = tqdm(total=len(chunks), desc="Transcribiendo", unit="chunk") if cfg.show_progress else None
    print("Esto es el chunk_audio: ", len(chunks))
    print("Esto es el chunk_audio: ", chunks[0])
    for idx, (s_ms, e_ms, chunk_audio) in enumerate(chunks, start=1):
        offset_sec = s_ms / 1000.0

        # Export temporal
        tmp_dir = str(Path(audio_path).parent)
        Path(tmp_dir).mkdir(parents=True, exist_ok=True)

        if cfg.export_mp3:
            suffix = f".chunk{idx}.mp3"
            print("Esto es el sufijo: ", suffix)
            print("Esto es el audio path: ", audio_path)
            print("Esto es el atmp_dir: ", tmp_dir)
            print("Esto es el chunk_audio: ", chunk_audio)
            with tempfile.NamedTemporaryFile(dir=tmp_dir, suffix=suffix, delete=False) as tmp:
                print(tmp.name)          
                chunk_audio.export(
                    tmp.name, format="mp3",
                    bitrate=cfg.mp3_bitrate,
                    parameters=["-ac", cfg.mp3_ac, "-ar", cfg.mp3_ar]
                )
                media_tmp = tmp.name
        else:
            suffix = f".chunk{idx}.wav"
            with tempfile.NamedTemporaryFile(dir=tmp_dir, suffix=suffix, delete=False) as tmp:
                chunk_audio.export(tmp.name, format="wav")
                media_tmp = tmp.name

        # Llamada a OpenAI (segments si es posible)
        try:
            with open(media_tmp, "rb") as f:
                tr = client.audio.transcriptions.create(
                    model=cfg.model,
                    file=f,
                    language=cfg.force_language or None,
                    response_format="verbose_json",
                    timestamp_granularities=["segment"],
                )
        except Exception:
            with open(media_tmp, "rb") as f:
                tr = client.audio.transcriptions.create(
                    model=cfg.model,
                    file=f,
                    language=cfg.force_language or None,
                    response_format="verbose_json",
                )

        # Construcción de párrafos ~15–22 s
        segments = getattr(tr, "segments", []) or []
        if segments:
            for seg in segments:
                s0 = float(_seg_val(seg, "start", 0.0) or 0.0)
                e0 = float(_seg_val(seg, "end", 0.0) or 0.0)
                text0 = (_seg_val(seg, "text", "") or "").strip()
                if not text0:
                    continue

                seg_start = offset_sec + s0
                seg_end = offset_sec + e0

                if not state["current_text"]:
                    state["current_text"] = text0
                    state["block_start"] = seg_start
                    state["last_end"] = seg_end
                else:
                    new_duration = seg_end - state["block_start"]
                    if new_duration <= (cfg.target_block_sec + cfg.max_slack_sec):
                        state["current_text"] = _merge_text_no_dupe(state["current_text"], text0)
                        state["last_end"] = seg_end
                    else:
                        flush_block()
                        state["current_text"] = text0
                        state["block_start"] = seg_start
                        state["last_end"] = seg_end
        else:
            # Fallback: usar texto completo
            t_all = (getattr(tr, "text", "") or "").strip()
            if t_all:
                if not state["current_text"]:
                    state["current_text"] = t_all
                    state["block_start"] = offset_sec
                    state["last_end"] = offset_sec
                else:
                    state["current_text"] = _merge_text_no_dupe(state["current_text"], t_all)
                    state["last_end"] = offset_sec

        # Limpieza
        try:
            os.remove(media_tmp)
        except Exception:
            pass

        if progress is not None:
            pct = (e_ms / duration_ms) * 100.0
            progress.set_postfix({"avance_audio": f"{pct:.1f}%"})
            progress.update(1)

    if progress is not None:
        progress.close()

    # Flush final
    flush_block()
    # Guardado final
    doc.save(output_docx)


    acta_docx, _ = build_acta_from_transcription(
    transcripcion_docx_path=output_docx,
    acta_file_path=output_docx,
    client=client,
    model="gpt-5",
    )
    print(f"✅ Listo. Word guardado en: {output_docx}")

    return output_docx, state["blocks_count"], duration_ms










